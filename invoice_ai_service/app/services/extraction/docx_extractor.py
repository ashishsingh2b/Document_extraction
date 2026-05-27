"""DOCX text extraction."""

import logging
from typing import Dict
import io
from docx import Document

logger = logging.getLogger(__name__)


class DocxExtractor:
    """Extract text from DOCX files."""
    
    def __init__(self):
        """Initialize DOCX extractor."""
        pass
    
    def extract_text(self, docx_data: bytes) -> Dict[str, any]:
        """
        Extract text from DOCX file.
        
        Args:
            docx_data: DOCX file as bytes
            
        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            # Open DOCX
            doc = Document(io.BytesIO(docx_data))
            
            # Extract paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = '\n'.join(text_parts)
            
            logger.info(f"Extracted {len(full_text)} characters from DOCX ({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables)")
            
            return {
                "text": full_text,
                "extraction_method": "docx",
                "text_length": len(full_text),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "success": True
            }
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            return {
                "text": "",
                "extraction_method": "failed",
                "error": str(e),
                "success": False
            }


# Global instance
docx_extractor = DocxExtractor()
